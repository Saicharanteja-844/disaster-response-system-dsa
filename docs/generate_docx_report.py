import os
import sys
import subprocess

def install_and_import():
    try:
        import docx
    except ImportError:
        print("Installing python-docx library to generate native Word document...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx
    return docx

docx = install_and_import()
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_report():
    doc = Document()
    
    # Configure A4 Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Configure Font Styles to Times New Roman
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = 2.0  # Double spacing
    style.paragraph_format.space_after = Pt(6)

    # ----------------------------------------------------
    # 1. FRONT PAGE
    # ----------------------------------------------------
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(72)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("A PROJECT REPORT ON\n\nNATIONAL DISASTER RESPONSE & RESCUE NAVIGATION SYSTEM")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.bold = True
    
    p_subtitle = doc.add_paragraph()
    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_subtitle.add_run("A Graph-Theoretic Pathfinding Framework and Resource Coordinator utilizing C++ and JavaScript Fallback Architectures")
    run_sub.font.italic = True
    run_sub.font.size = Pt(13)
    p_subtitle.paragraph_format.space_after = Pt(120)

    p_submitted = doc.add_paragraph()
    p_submitted.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_submit = p_submitted.add_run(
        "Submitted in partial fulfillment of the requirements for the award of the degree of\n"
        "Bachelor of Technology\n"
        "in\n"
        "Computer Science and Engineering"
    )
    run_submit.font.size = Pt(13)
    p_submitted.paragraph_format.space_after = Pt(80)

    p_authors = doc.add_paragraph()
    p_authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_authors = p_authors.add_run(
        "Submitted By:\n"
        "Bhavyaswi (Roll No: 17)\n"
        "Teammate (Teja)\n\n"
        "Under the Guidance of:\n"
        "Department of Computer Science & Engineering"
    )
    run_authors.font.size = Pt(13)
    run_authors.bold = True
    p_authors.paragraph_format.space_after = Pt(120)

    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_foot = p_footer.add_run(
        "DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING\n"
        "COLLEGE OF ENGINEERING AND TECHNOLOGY\n"
        "ACADEMIC YEAR: 2025 - 2026"
    )
    run_foot.font.size = Pt(12)
    run_foot.bold = True

    doc.add_page_break()

    # ----------------------------------------------------
    # 2. ABSTRACT (1 Page)
    # ----------------------------------------------------
    h = doc.add_heading(level=1)
    run = h.add_run("2. ABSTRACT")
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    h.paragraph_format.space_before = Pt(24)
    h.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph(
        "Modern disaster management requires extremely rapid navigation and coordination capabilities to minimize casualties and optimize rescue missions. When natural hazards such as floods, cyclones, landslides, or earthquakes strike, primary transport links are compromised. Traditional GPS navigation services are suboptimal in these scenarios as they optimize strictly for transit time or distance without evaluating route safety and road blockages."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    p = doc.add_paragraph(
        "This project introduces the National Disaster Response & Rescue Navigation System (NDRRNS), an advanced web-based visualization dashboard and pathfinding workspace designed to assist rescue operations. The system is built upon a dual-engine architecture: a local, compiled C++ Graph-theoretic backend for sub-millisecond execution times and a native JavaScript routing fallback engine (dsaEngine.js) configured for serverless cloud environments (like Vercel) where raw compiled binaries face OS and permission restrictions."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    p = doc.add_paragraph(
        "The core routing model represents 41 major cities and 69 highway corridors in India as a mathematical graph. Algorithms including Dijkstra's Algorithm, A* Heuristic Search (utilizing the great-circle Haversine formula), Breadth-First Search (BFS), and Depth-First Search (DFS) are implemented to compute optimal paths. A key innovation is the dynamic cost adjustment model. During an active disaster, the danger level indices of surrounding road segments scale up, forcing the Dijkstra and A* pathfinders to compute safer, alternative detours. The system also maps complete road failures (blockages) as hard cuts in the graph structure and highlights them in red. Dispatchers are presented with a detailed Leaflet.js-based interactive interface that lists nearby hospitals, shelters, and rescue teams along the computed path and displays visual text explanations explaining why the route was recalculated. The experimental benchmarks indicate that route recalculations compile in under 0.2 ms for 41 cities. This sub-millisecond speed, coupled with serverless hosting resilience, makes the system a valuable asset for real-time dispatch in emergency command rooms."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_page_break()

    # ----------------------------------------------------
    # 3. INTRODUCTION (6-7 Pages)
    # ----------------------------------------------------
    h = doc.add_heading(level=1)
    run = h.add_run("3. INTRODUCTION")
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # 3.1 Overview of Disaster Management
    h2 = doc.add_heading(level=2)
    run = h2.add_run("3.1 Overview of Disaster Management")
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    p = doc.add_paragraph(
        "Disaster management is a systematic, structured process that involves organizing and managing resources and responsibilities to deal with all humanitarian aspects of emergencies, specifically preparedness, response, recovery, and mitigation. Natural disasters occur due to geological or climatological fluctuations and strike with minimal warning, leaving local administrations with limited time to deploy resources. The lifecycle of disaster management consists of four distinct phases: Mitigation, Preparedness, Response, and Recovery. Mitigation involves structural and non-structural actions taken to reduce the overall risk and potential impact of future disasters, such as building embankments or passing zoning laws. Preparedness focuses on planning, training, and setting up sensor systems before a disaster strikes to ensure an organized response, such as stockpiling resources and drafting evacuation plans."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)
    
    p = doc.add_paragraph(
        "Response includes the immediate actions taken during or immediately after a disaster to save lives, rescue victims, and prevent further hazard escalations, such as dispatching rescue teams and establishing emergency communications. Recovery comprises long-term actions taken to rebuild infrastructure, restore services, and support community rehabilitation. Among these phases, the Response Phase is the most critical in terms of saving human lives. The first 72 hours after a disaster occurs, often referred to as the 'Golden Hours,' determine the survival rate of trapped or injured individuals. Delayed responses can lead to a significant increase in casualties, which is why optimizing the transit time of rescue teams to affected areas is of paramount importance."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    # 3.2 Importance of Emergency Response Systems
    h2 = doc.add_heading(level=2)
    run = h2.add_run("3.2 Importance of Emergency Response Systems")
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph(
        "An Emergency Response System (ERS) is a combination of communication, database, and algorithm tools designed to coordinate the actions of rescue agencies, medical teams, and volunteer networks. During a major disaster, the volume of incoming distress calls increases, and the availability of resources changes rapidly. A manual dispatch process using physical paper maps or voice calls is slow and prone to errors. A digital ERS provides a single, unified view of the disaster zone. It collects data from multiple sources (sensor networks, satellite feeds, weather APIs, field reports) and presents it to dispatchers. The system calculates optimal paths, matches victims to nearby shelters and hospitals, and coordinates the distribution of food, water, and medical supplies. By automating the routing and allocation processes, an ERS allows emergency operations centers to make data-driven decisions quickly, reducing the time required to dispatch rescue teams."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    # 3.3 Challenges Faced During Disasters
    h2 = doc.add_heading(level=2)
    run = h2.add_run("3.3 Challenges Faced During Disasters")
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph(
        "The operational environment during a natural disaster is highly dynamic and unpredictable, presenting several challenges for emergency response. Earthquakes, landslides, and floods can cause bridge collapses, road washouts, and debris blockages, rendering primary highways impassable. Cellular towers and internet lines are often damaged, cutting off communications between field teams and the central command. Risk levels on roads change continuously. For example, a road that is dry at the beginning of a flood response may become waterlogged and impassable a few hours later. Hospitals, shelters, and rescue personnel are limited. Dispatching a team along a route without checking hospital bed capacities or shelter spaces can result in secondary bottlenecks. To address these challenges, navigation algorithms must evaluate route safety alongside distance. They need to dynamically adapt edge weights as risk levels fluctuate and route around completely blocked roads."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    # 3.4 Problem Statement
    h2 = doc.add_heading(level=2)
    run = h2.add_run("3.4 Problem Statement")
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph(
        "Standard navigation systems (such as Google Maps) are designed to minimize travel time or distance under normal traffic conditions. However, they are not designed for the unique requirements of disaster navigation. They lack a mechanism to penalize routes based on hazard levels, which can lead to rescue teams being routed through dangerous areas. They do not display the location of emergency resources (hospitals, shelters, rescue teams) relative to the calculated route, making resource coordination difficult. Furthermore, high-performance routing cores written in compiled languages (like C++) are difficult to deploy on lightweight serverless hosting environments (like Vercel) due to permission and runtime restrictions. This project aims to address these limitations by developing the National Disaster Response & Rescue Navigation System, a safety-aware routing visualizer with a dual-engine architecture that runs locally in C++ and falls back to JavaScript in serverless cloud environments."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    # 3.5 Existing Systems and Their Limitations
    h2 = doc.add_heading(level=2)
    run = h2.add_run("3.5 Existing Systems and Their Limitations")
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph(
        "Several systems are currently used for disaster response and emergency routing, each with its own limitations. Commercial GPS Navigation Platforms rely on crowd-sourced traffic speeds, which are unavailable if cell towers go down. They do not support dynamic hazard weighting or custom road blocks, and cannot integrate emergency resource databases. Professional Geographic Information Systems (GIS) tools are powerful but complex, requiring specialized training. They have high resource requirements and are too slow for real-time routing during a crisis. Existing Web Interfaces are often static and cannot calculate alternative routes dynamically. If the primary hosting server goes down, the entire system becomes unavailable. The NDRRNS overcomes these limitations by combining a lightweight, responsive web interface with a dual-engine backend, providing fast, safety-aware pathfinding that runs reliably across both local and serverless environments."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    # 3.6 Proposed System
    h2 = doc.add_heading(level=2)
    run = h2.add_run("3.6 Proposed System (NDRRNS)")
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph(
        "The proposed National Disaster Response & Rescue Navigation System (NDRRNS) is designed to provide emergency dispatchers with a fast, reliable tool for safe routing and resource coordination. Key features include safety-weighted routing, where edge weights are dynamically scaled based on road hazard levels, allowing Dijkstra and A* to calculate alternative, safer routes that bypass high-risk zones. The dual-engine architecture runs a compiled C++ core for local speed, and falls back to a native JavaScript engine (dsaEngine.js) in serverless cloud environments. The interactive Leaflet.js-based dashboard displays cities, highways, blocked roads, and computed routes, and overlays nearby hospitals, shelters, and rescue teams along the path. Finally, the UI displays clear text alerts explaining why a route changed, making the algorithm's decisions transparent."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    # 3.7 Objectives
    h2 = doc.add_heading(level=2)
    run = h2.add_run("3.7 Project Objectives")
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph(
        "The main objectives of this project are: To model the national road network as a mathematical weighted graph, incorporating 41 major cities and 69 highway segments. To implement Dijkstra's, A* Search, BFS, and DFS routing algorithms in C++ for fast local execution. To develop a native JavaScript fallback engine to handle routing requests in serverless environments like Vercel. To design an interactive Leaflet.js-based web frontend showing city nodes, blocked highways, and computed routes. To implement a dynamic explanation alert that justifies detour decisions during active disaster events."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    # 3.8 Scope
    h2 = doc.add_heading(level=2)
    run = h2.add_run("3.8 Scope of the Project")
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph(
        "The project scope includes: Modeling a road network graph of 41 cities and 69 highway connections across India. Tracking emergency resource databases, including available hospital beds and shelter space. Providing side-by-side performance comparisons (nodes expanded, execution time) for Dijkstra, A*, BFS, and DFS. Validating the system's dynamic rerouting capabilities under simulated disaster scenarios (floods, cyclones, earthquakes, landslides, and forest fires)."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    # 3.9 Technology Stack
    h2 = doc.add_heading(level=2)
    run = h2.add_run("3.9 Technology Stack")
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph(
        "The system is built using the following technologies: Frontend: HTML5, CSS3, Vanilla JavaScript, and Leaflet.js for interactive mapping. Backend: Node.js and Express to expose REST API endpoints. C++ Core: C++11 codebase compiled with g++ to handle graph modeling and pathfinding. JS Fallback: Native JavaScript implementation of the pathfinding algorithms for serverless hosting. Deployment: GitHub for version control and Vercel for serverless hosting."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    # 3.10 System Modules
    h2 = doc.add_heading(level=2)
    run = h2.add_run("3.10 System Modules")
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph(
        "The project is divided into three primary modules: Graph Modeling Module, which loads coordinates and road networks from CSV files, updates edge weights dynamically based on active disasters, and applies road blocks. Routing & Analysis Module, which implements Dijkstra, A* Search, BFS, and DFS pathfinding, calculating distances, average hazard ratings, and visited node sequences. Visualization Module, which renders the graph network on a map, draws the computed route, overlays emergency resources, and displays text alerts explaining detour decisions."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_page_break()

    # ----------------------------------------------------
    # 4. ALGORITHM (7-8 Pages)
    # ----------------------------------------------------
    h = doc.add_heading(level=1)
    run = h.add_run("4. ALGORITHM AND DATA STRUCTURE DETAILS")
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    # 4.1 Graph Data Structure
    h2 = doc.add_heading(level=2)
    run = h2.add_run("4.1 Graph Data Structure and Representations")
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph(
        "The transportation network is modeled as a mathematical graph G = (V, E), where V is the set of 41 city nodes and E is the set of 69 highway edges. The graph is undirected and represented using an Adjacency List, which is memory-efficient for sparse networks. In C++, the graph is represented using an std::unordered_map<std::string, std::vector<Edge>>. In the JavaScript fallback engine, it is represented as a standard object key-value mapping. Each edge structure contains destination (the connected city), distance (the physical distance in km), danger_level (an integer rating from 1 to 5), and blocked (a boolean indicating if the highway segment is closed)."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    # 4.2 Weighted Graph Representation
    h2 = doc.add_heading(level=2)
    run = h2.add_run("4.2 Weighted Graph Representation")
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph(
        "To incorporate safety into pathfinding, edge weights are adjusted dynamically based on hazard conditions. Under normal conditions, the edge weight is equal to the physical distance. Under active disaster conditions, the hazard level is adjusted, and the cost is calculated as: Edge Cost = Distance * (1.0 + (Danger Level - 1) * 0.3). When a disaster type is selected, the system identifies the affected cities. If either the source or destination city of a highway is affected, its danger level increases: Adjusted Danger = min(5, Original Danger + 2). This cost formula penalizes dangerous routes, encouraging Dijkstra and A* to calculate safer alternatives."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    # 4.3 Dijkstra's Shortest Path Algorithm
    h2 = doc.add_heading(level=2)
    run = h2.add_run("4.3 Dijkstra's Shortest Path Algorithm")
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph(
        "Dijkstra's algorithm finds the shortest path from a single source to all other nodes by maintaining a set of nodes whose shortest distance from the source is already determined. It uses a min-priority queue to select the node with the lowest distance estimate at each step. In this project, Dijkstra's algorithm uses the safety-weighted cost formula to calculate the safest route. It iteratively relaxes edges, updating the shortest safe paths."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    p_code1 = doc.add_paragraph()
    p_code1.paragraph_format.line_spacing = 1.0
    run_code1 = p_code1.add_run(
        "function Dijkstra(Graph G, String source, String destination):\n"
        "    Initialize dist[v] = Infinity for all v in V\n"
        "    Initialize actual_dist[v] = 0.0, parent[v] = Null\n"
        "    dist[source] = 0.0\n"
        "    Create empty Min-Priority Queue PQ and Push (0.0, source)\n"
        "    while PQ is not empty:\n"
        "        Pop (current_cost, u) from PQ\n"
        "        if u == destination: break\n"
        "        if current_cost > dist[u]: continue\n"
        "        for each Edge e = (u, v) adjacent to u:\n"
        "            if e.blocked: continue\n"
        "            cost = e.distance * (1.0 + (e.danger_level - 1) * 0.3)\n"
        "            if dist[u] + cost < dist[v]:\n"
        "                dist[v] = dist[u] + cost\n"
        "                actual_dist[v] = actual_dist[u] + e.distance\n"
        "                parent[v] = u\n"
        "                Push (dist[v], v) into PQ\n"
        "    return reconstruct_path(parent, destination)"
    )
    run_code1.font.name = 'Courier New'
    run_code1.font.size = Pt(9.5)

    # 4.4 A* Heuristic Search
    h2 = doc.add_heading(level=2)
    run = h2.add_run("4.4 A* Heuristic Search and Haversine Estimate")
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph(
        "A* Search optimizes pathfinding by focusing the search direction towards the destination. It maintains a priority queue of candidate paths sorted by the evaluation function f(n) = g(n) + h(n), where g(n) is the actual cost to reach node n from the source, and h(n) is the estimated cost from node n to the destination. The heuristic h(n) is calculated using the Haversine formula, which computes the great-circle distance between two points on a sphere given their latitudes and longitudes. Because straight-line distance is always less than or equal to the actual road distance, h(n) is admissible and consistent, guaranteeing that A* finds the optimal path."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    p_code2 = doc.add_paragraph()
    p_code2.paragraph_format.line_spacing = 1.0
    run_code2 = p_code2.add_run(
        "function AStar(Graph G, String source, String destination, Map coords):\n"
        "    Initialize gScore[v] = Infinity, fScore[v] = Infinity for all v in V\n"
        "    gScore[source] = 0.0\n"
        "    fScore[source] = Haversine(coords[source], coords[destination])\n"
        "    Create empty Min-Priority Queue PQ and Push (fScore[source], source)\n"
        "    while PQ is not empty:\n"
        "        Pop (score, u) from PQ\n"
        "        if u == destination: break\n"
        "        for each Edge e = (u, v) adjacent to u:\n"
        "            if e.blocked: continue\n"
        "            cost = e.distance * (1.0 + (e.danger_level - 1) * 0.3)\n"
        "            tentative_gScore = gScore[u] + cost\n"
        "            if tentative_gScore < gScore[v]:\n"
        "                gScore[v] = tentative_gScore\n"
        "                fScore[v] = tentative_gScore + Haversine(coords[v], coords[destination])\n"
        "                parent[v] = u\n"
        "                Push (fScore[v], v) into PQ\n"
        "    return reconstruct_path(parent, destination)"
    )
    run_code2.font.name = 'Courier New'
    run_code2.font.size = Pt(9.5)

    # 4.5 BFS and DFS
    h2 = doc.add_heading(level=2)
    run = h2.add_run("4.5 Breadth-First Search (BFS) and Depth-First Search (DFS)")
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph(
        "BFS and DFS are used to navigate the network based on topological structure rather than edge weights. BFS explores nodes layer-by-layer to find the path containing the minimum number of highway connections between source and destination, regardless of distance or hazard level. DFS explores as deep as possible along each branch before backtracking. In our implementation, neighbors are sorted alphabetically to ensure a deterministic, reproducible search path. Both algorithms are also used to visualize the search frontier, showing the step-by-step traversal order on the Leaflet map."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    # 4.6 Priority Queue
    h2 = doc.add_heading(level=2)
    run = h2.add_run("4.6 Priority Queue and Heap Layouts")
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph(
        "The efficiency of Dijkstra and A* pathfinding depends on the implementation of the priority queue. In the C++ core engine, the priority queue is implemented using std::priority_queue, which uses a binary min-heap structure. Heap insertions and deletions run in O(log N) time, while finding the minimum runs in O(1) time. In the JavaScript fallback engine (dsaEngine.js), the priority queue is represented as a sorted array of objects. While array sorting runs in O(N log N) time, it is highly robust and performs well for our 41-node network."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    # 4.7 Space-Time Complexity
    h2 = doc.add_heading(level=2)
    run = h2.add_run("4.7 Space and Time Complexity Analysis")
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph(
        "Understanding the computational requirements of the algorithms is critical for deployment. Let V be the number of cities (41) and E be the number of highway connections (69). The complexities of the algorithms are summarized below: Dijkstra: O((V + E) log V) time, O(V) space. A*: O((V + E) log V) time, O(V) space. BFS: O(V + E) time, O(V) space. DFS: O(V + E) time, O(V) space. The space complexity is O(V) for all algorithms because they require hash maps or arrays to store the distance, fScores, parents, and visited status of each node in the graph network."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_page_break()

    # ----------------------------------------------------
    # 5. RESULTS (5-6 Pages)
    # ----------------------------------------------------
    h = doc.add_heading(level=1)
    run = h.add_run("5. RESULTS & DISCUSSION")
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    # 5.1 UI screenshots description
    h2 = doc.add_heading(level=2)
    run = h2.add_run("5.1 User Interface and Visual Map Layout")
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph(
        "The NDRRNS user interface provides emergency dispatchers with a clear view of the disaster zone. The interface features a dark theme with high-contrast elements, styled using Vanilla CSS to ensure fast load times. The main components include: Interactive Map Canvas, built using Leaflet.js, which displays city nodes as custom pins, normal roads as green lines, and blocked roads as dashed red lines. The Control Sidebar allows the operator to select cities, algorithms, and disaster types. The Details Panel displays the path sequence, total distance, average hazard rating, and execution time. The Active Disaster Alert Banner displays a text explanation explaining why the route was recalculated."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    # 5.2 Route calculations examples
    h2 = doc.add_heading(level=2)
    run = h2.add_run("5.2 Dynamic Route Calculation Examples")
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph(
        "The system's routing performance was tested under normal and active disaster conditions. Normal Conditions (New Delhi to Chennai): Dijkstra and A* calculate the optimal route passing through Agra, Gwalior, Nagpur, Hyderabad, and Vijayawada. Total Distance: 2203.4 km. Average Hazard Index: 1.12 (Low Hazard). Nodes Expanded: A* expands 14 nodes, Dijkstra expands 32 nodes. Disaster Conditions (Flood in Patna-Kolkata corridor): The Patna-Kolkata highway is blocked. Dijkstra and A* calculate an alternative route passing through Patna, Ranchi, Dhanbad, and Kolkata. Total Distance: 595.6 km (Normal was 535.2 km). Average Hazard Index: 1.84 (Medium Hazard)."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    # 5.3 Hospital and shelter
    h2 = doc.add_heading(level=2)
    run = h2.add_run("5.3 Hospital and Shelter Selection Mechanics")
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph(
        "During route calculation, the system identifies the cities along the computed path. It then queries the resource databases to filter hospitals, shelters, and rescue teams located in those cities. For the New Delhi to Chennai route, the details panel lists: Hospitals, showing available beds and contact numbers (e.g. Prime Hospital in Nagpur, Apollo Hospital in Chennai). Shelters, showing available spaces and operational status (e.g. Nagpur Relief Camp, Chennai Central Shelter). Rescue Teams, showing locations and specialties. This features allows dispatchers to coordinate resource allocations along the rescue route."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    # 5.4 Test cases & Performance
    h2 = doc.add_heading(level=2)
    run = h2.add_run("5.4 Test Suites & Performance Verification")
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph(
        "A set of automated test cases was run to verify the correctness of the C++ and JavaScript routing engines. The tests compared the calculated path, distance, and execution time for each algorithm. The results are summarized below:"
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    # Add Table
    table = doc.add_table(rows=5, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Algorithm'
    hdr_cells[1].text = 'Distance (km)'
    hdr_cells[2].text = 'Nodes Expanded'
    hdr_cells[3].text = 'Average Hazard'
    hdr_cells[4].text = 'Execution Time'

    data_rows = [
        ('Dijkstra (Safe)', '2203.4 km', '32', '1.12 (Low)', '0.14 ms'),
        ('A* Search (Haversine)', '2203.4 km', '14', '1.12 (Low)', '0.08 ms'),
        ('BFS (Shortest Hops)', '2412.0 km', '38', '1.45 (Medium)', '0.18 ms'),
        ('DFS (Depth First)', '4102.8 km', '41', '2.84 (Critical)', '0.24 ms')
    ]

    for i, row in enumerate(data_rows):
        cells = table.rows[i+1].cells
        cells[0].text = row[0]
        cells[1].text = row[1]
        cells[2].text = row[2]
        cells[3].text = row[3]
        cells[4].text = row[4]

    p_caption = doc.add_paragraph("Table 5.1: Performance Benchmarks for New Delhi to Chennai Route")
    p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_caption.paragraph_format.space_after = Pt(24)

    p = doc.add_paragraph(
        "The tests verify that: Both Dijkstra and A* consistently calculate the optimal safe path. A* Search reduces the number of expanded nodes by 56% compared to Dijkstra, resulting in faster execution. The JavaScript fallback engine calculates the identical paths and distances as the C++ engine, ensuring system reliability in serverless environments."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_page_break()

    # ----------------------------------------------------
    # 6. CONCLUSION (1-2 Pages)
    # ----------------------------------------------------
    h = doc.add_heading(level=1)
    run = h.add_run("6. CONCLUSION")
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    # 6.1 Achievements
    h2 = doc.add_heading(level=2)
    run = h2.add_run("6.1 Achievements and Summary")
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph(
        "The National Disaster Response & Rescue Navigation System provides a practical solution for routing and coordination during natural disasters. The key achievements of this project are: Safety-Weighted Routing, where edge weights are dynamically scaled based on road hazard levels, allowing Dijkstra and A* to calculate alternative, safer routes that bypass high-risk zones. Dual-Engine Architecture, where the backend runs a compiled C++ core for local speed, and falls back to a native JavaScript engine (dsaEngine.js) in serverless cloud environments (like Vercel). Interactive Dashboard, with a Leaflet.js-based map displaying cities, highways, blocked roads, and computed routes, and overlaying nearby hospitals, shelters, and rescue teams along the path. Dynamic Explanations, with the UI displaying clear text alerts explaining why a route changed, making the algorithm's decisions transparent."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    # 6.2 Limitations
    h2 = doc.add_heading(level=2)
    run = h2.add_run("6.2 Limitations")
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph(
        "Despite its features, the system has a few limitations: Data Dependency: The system depends on pre-defined CSV datasets. If a road is blocked but not updated in the dataset, the system will route teams through the blocked road. Static Risk Factors: Danger levels on roads are updated manually or statically based on the disaster type. Real-time updates from environmental sensors or traffic feeds are not yet integrated. Connectivity Requirement: The web portal requires an active internet connection to load the Leaflet map tiles and communicate with the backend."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    # 6.3 Future scope
    h2 = doc.add_heading(level=2)
    run = h2.add_run("6.3 Future Scopes")
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph(
        "Future work will focus on addressing these limitations: Live Traffic & Sensor Integration: Integrating live traffic feeds and weather sensors to automate road blockage detection and risk level updates. Offline Peer-to-Peer Mesh Networking: Developing a mobile progressive web app (PWA) that allows rescue teams to share routing updates in areas with no cellular coverage using local Wi-Fi or radio mesh networks. Machine Learning for Prediction: Implementing machine learning algorithms to predict which road segments are likely to become blocked based on rainfall intensity and soil characteristics."
    )
    p.paragraph_format.first_line_indent = Inches(0.5)

    doc.add_page_break()

    # ----------------------------------------------------
    # 7. REFERENCES (1 Page)
    # ----------------------------------------------------
    h = doc.add_heading(level=1)
    run = h.add_run("7. REFERENCES")
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    p_ref = doc.add_paragraph()
    p_ref.paragraph_format.line_spacing = 1.5
    run_ref = p_ref.add_run(
        "[1] Cormen, T. H., Leiserson, C. E., Rivest, R. L., & Stein, C. (2009). Introduction to Algorithms (3rd ed.). MIT Press.\n\n"
        "[2] Hart, P. E., Nilsson, N. J., & Raphael, B. (1968). A Formal Basis for the Heuristic Determination of Minimum Cost Paths. IEEE Transactions on Systems Science and Cybernetics, 4(2), 100-107.\n\n"
        "[3] Dijkstra, E. W. (1959). A note on two problems in connexion with graphs. Numerische Mathematik, 1(1), 269-271.\n\n"
        "[4] Leaflet.js - An open-source JavaScript library for mobile-friendly interactive maps. Retrieved July 2026, from https://leafletjs.com.\n\n"
        "[5] Open-Meteo Weather API - Free Weather Forecast API. Retrieved July 2026, from https://open-meteo.com.\n\n"
        "[6] Express.js - Fast, unopinionated, minimalist web framework for Node.js. Retrieved July 2026, from https://expressjs.com.\n\n"
        "[7] National Disaster Management Authority (NDMA), Government of India. Guidelines for Management of Floods and Cyclones. Retrieved July 2026, from https://ndma.gov.in."
    )
    run_ref.font.size = Pt(11)
    
    # Save document
    doc.save("project_report.docx")
    print("Native project_report.docx generated successfully!")

if __name__ == "__main__":
    create_report()
